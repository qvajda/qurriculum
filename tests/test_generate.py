import docx

from qurriculum import generate


def test_default_out_path():
    assert generate.DEFAULT_OUT == generate.Path("out/be-fr-analytics-fr.docx")


def test_docx_opens_without_error(tmp_path):
    out = tmp_path / "fr.docx"
    generate.render(generate.load_content(), out)
    docx.Document(out)


def test_sections_render_in_declared_order(tmp_path):
    out = tmp_path / "fr.docx"
    content = generate.load_content()
    generate.render(content, out)
    text = [p.text for p in docx.Document(out).paragraphs]
    positions = [text.index(s["heading"]) for s in content["sections"]]
    assert positions == sorted(positions), positions


def test_name_and_entries_are_present(tmp_path):
    out = tmp_path / "fr.docx"
    content = generate.load_content()
    generate.render(content, out)
    full_text = "\n".join(p.text for p in docx.Document(out).paragraphs)
    assert content["name"] in full_text
    for section in content["sections"]:
        for entry in section["entries"]:
            assert entry in full_text
