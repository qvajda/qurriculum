"""Renders the BE:FR analytics FR core as a .docx from committed YAML content.

Structure and prose both live in content/be_fr_analytics.yml (ADR-0001 D-1);
this module walks whatever sections it is given and carries no section names
or ordering of its own.
"""

import argparse
from pathlib import Path

import docx
import yaml

CONTENT_PATH = Path(__file__).parent / "content" / "be_fr_analytics.yml"
DEFAULT_OUT = Path("out/be-fr-analytics-fr.docx")


def load_content(path: Path = CONTENT_PATH) -> dict:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def render(content: dict, out: Path) -> Path:
    document = docx.Document()
    document.add_heading(content["name"], level=0)
    document.add_paragraph(content["title"])
    for line in content["contact"]:
        document.add_paragraph(line)
    for section in content["sections"]:
        document.add_heading(section["heading"], level=1)
        for entry in section["entries"]:
            document.add_paragraph(entry)
    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(out)
    return out


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    render(load_content(), args.out)


if __name__ == "__main__":
    main()
